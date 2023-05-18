import time
import configparser
import openai
import plotly.graph_objects as go
import plotly
import pandas as pd
import numpy as np
import json
import hashlib
import threading
import dash_bootstrap_components as dbc
from dash import html, dash_table
import os
import docx
from docx.shared import Pt
from docx.enum.style import WD_STYLE_TYPE
import logging


class ChatGpt:
    def __init__(self):
        openai.api_key = self.__read_api_key()
        self.persona = "Você é um analista de dados experiente. Graduado em Ciência da computação com"\
                       " pós-graduação em inteligência artificial, mestrado em matemática e doutorado em"\
                       " estatística"

    def __read_api_key(self):
        config = configparser.ConfigParser()
        config.read('C:\\Users\\luiz henrique\\Documents\\dash_and_chatgpt\\api_key_openai.ini')
        return config['OPENAI']['ApiKey']

    def ask(self, prompt, data):
        if data is not None:
            response = openai.ChatCompletion.create(
                model="gpt-3.5-turbo",
                temperature=0.4,
                messages=[
                    {"role": "system", "content": self.persona},
                    {"role": "user", "content": 'os dados que você vai utilizar:\n\n' + data},
                    {"role": "user", "content": prompt}
                ]
            )
        else:
            response = openai.ChatCompletion.create(
                model="gpt-3.5-turbo",
                temperature=0.9,
                messages=[
                    {"role": "user", "content": prompt}
                ]
            )
        return response.to_dict()['choices'][0].to_dict()['message']['content'].strip()


class Graphs:
    @staticmethod
    def Bar(x, y, text=None, height=350):
        fig = go.Figure()
        if text is not None:
            fig.add_trace(go.Bar(x=x, y=y, text=text, textposition='auto'))
        else:
            fig.add_trace(go.Bar(x=x, y=y))
        fig.update_layout(xaxis={'categoryorder': 'total descending'}, margin=dict(l=0, r=0, t=0, b=0), height=height)
        return fig

    @staticmethod
    def BarH(x, y, text=None, height=350):
        fig = go.Figure()
        if text is not None:
            fig.add_trace(go.Bar(x=x, y=y, text=text, textposition='auto', orientation='h'))
        else:
            fig.add_trace(go.Bar(x=x, y=y, orientation='h'))
        fig.update_layout(yaxis={'categoryorder': 'total ascending'}, margin=dict(l=0, r=0, t=0, b=0), height=height)
        return fig


def ask_chatgpt(prompt, data):
    chatgpt = ChatGpt()
    try:
        response = chatgpt.ask(prompt, data)
    except openai.InvalidRequestError as msg:
        return msg.user_message
    else:
        return response


def read_df_and_select_pais(data, active_tab):
    """
    Cria um data frame e caso a aba seja diferente de 'todos', faz uma seleção
    dos dados pelo pais e retorna o dataframe.
    """
    dframe = pd.read_json(
        data, orient='split', convert_dates=['Data Cadastro'],
        dtype={'Filhos em Casa': str, 'Adolescentes em Casa': str}
    )
    pais = active_tab.split('_')[1]
    if pais != 'todos':
        dframe = dframe.query(f"Pais == '{pais}'")
    return dframe


def save_narrative(key, narrative):
    try:
        path_arquivo = 'C:\\Users\\luiz henrique\\Documents\\dash_and_chatgpt\\historico_narrativas.json'
        with open(path_arquivo, 'r') as json_file:
            historico = json.load(json_file)
    except (json.JSONDecodeError, FileNotFoundError):
        logging.exception('Não foi possivel abrir o arquivo com o historico das narrativas')
    else:
        # é feita uma conversão do texto em hash e essa hash é comparada com a hash de outros textos
        # dentro do arquivo para saber se já existe um texto parecido dentro do arquivo caso exista
        # o texto não é salvo, caso contrário o texto e seu hash são salvos no arquivo.
        if hashlib.md5(narrative.encode()).hexdigest() not in historico[key]['Hashs']:
            if len(historico[key]['Hashs']) == 5:
                del historico[key]['Hashs'][0]
                del historico[key]['Values'][0]
                historico[key]['Hashs'].append(hashlib.md5(narrative.encode()).hexdigest())
                historico[key]['Values'].append(narrative)
            else:
                historico[key]['Hashs'].append(hashlib.md5(narrative.encode()).hexdigest())
                historico[key]['Values'].append(narrative)
        try:
            with open(path_arquivo, 'w') as json_file:
                json.dump(historico, json_file)
        except (json.JSONDecodeError, FileNotFoundError):
            logging.exception('Não foi possivel salvar o historico')


def response_modal_prompt_narrative(df, json_key):
    try:
        with open('C:\\Users\\luiz henrique\\Documents\\dash_and_chatgpt\\prompt_chatGpt.json', 'r') as json_file:
            prompt_json = json.load(json_file)
    except (json.JSONDecodeError, FileNotFoundError):
        raise Exception('Não foi possivel abrir o arquivo!')
    else:
        return (
            True, prompt_json[json_key],
            dash_table.DataTable(
                data=df.to_dict(orient='records'),
                columns=[{'id': c, 'name': c} for c in df.columns],
                page_action='none',
                style_table={'height': '200px', 'overflowY': 'auto'},
                style_cell={'textAlign': 'left', 'fontWeight': 'bold'}
            )
        )


def save_json_prompt(json_key, value):
    try:
        with open('C:\\Users\\luiz henrique\\Documents\\dash_and_chatgpt\\prompt_chatGpt.json', 'r') as json_file:
            prompt_json = json.load(json_file)

        prompt_json[json_key] = value

        with open('C:\\Users\\luiz henrique\\Documents\\dash_and_chatgpt\\prompt_chatGpt.json', 'w') as json_file:
            json.dump(prompt_json, json_file)
    except (json.JSONDecodeError, FileNotFoundError):
        logging.exception('Não foi possivel abrir o arquivo!')
        return dbc.Alert(
            'Não foi possivel salvar o prompt!', id="alert-auto", is_open=True, duration=2000, color='danger'
        )
    else:
        return dbc.Alert(
            'Prompt salvo!', id="alert-auto", is_open=True, duration=2000, color='success'
        )


def modal_historic_narrative(json_key):
    try:
        path_arquivo = 'C:\\Users\\luiz henrique\\Documents\\dash_and_chatgpt\\historico_narrativas.json'
        with open(path_arquivo, 'r') as json_file:
            historico = json.load(json_file)
    except (json.JSONDecodeError, FileNotFoundError):
        raise Exception('Não foi possivel abrir o arquivo!')
    else:
        if len(historico[json_key]['Values']) > 0:
            return True, dbc.Accordion(
                [
                    dbc.AccordionItem(
                        [
                            dbc.Textarea(value=t.strip(), className="mb-3", size='md', rows=10)
                        ], title=t.strip()[:100] + '...'
                    )
                    for t in historico[json_key]['Values']
                ],
                start_collapsed=True
            )
        else:
            return True, html.Label('o Historico esta vazio!!')


def create_narrative(df: str, json_key: str) -> str:
    try:
        with open('C:\\Users\\luiz henrique\\Documents\\dash_and_chatgpt\\prompt_chatGpt.json', 'r') as json_file:
            prompt_json = json.load(json_file)
    except (json.JSONDecodeError, FileNotFoundError):
        return "Houve um problema ao ler o arquivo que armazena o prompt!"
    else:
        prompt = prompt_json[json_key]

        narrative = ask_chatgpt(prompt, df)

        # Salva o retorno do ChatGpt no arquivo de histórico de narrativas
        threading.Thread(target=save_narrative, args=[json_key, narrative]).start()
        return narrative


def create_false_narrative(json_key):
    try:
        path_arquivo = 'C:\\Users\\luiz henrique\\Documents\\dash_and_chatgpt\\historico_narrativas.json'
        with open(path_arquivo, 'r') as json_file:
            historico = json.load(json_file)
    except (json.JSONDecodeError, FileNotFoundError):
        raise Exception('Não foi possivel abrir o arquivo!')
    else:
        size = len(historico[json_key]['Values'])
        if size > 1:
            time.sleep(2)
            return historico[json_key]['Values'][np.random.randint(low=0, high=size-1)]
        else:
            time.sleep(2)
            return historico[json_key]['Values'][0]


def relatorio(json_figures, json_relatorios):
    figures = json.loads(json_figures)
    relatorios = json.loads(json_relatorios)

    for k1 in figures.keys():
        for k2 in figures[k1].keys():
            if k2 != 'Campanha Table' and figures[k1][k2]:
                plotly.io.write_image(
                    figures[k1][k2],
                    os.path.join('.', 'images', k2.lower().replace(' ', '_') + '.png'), 'png'
                )

    doc = docx.Document()

    styles_element = doc.styles.element
    rpr_default = styles_element.xpath('./w:docDefaults/w:rPrDefault/w:rPr')[0]
    lang_default = rpr_default.xpath('w:lang')[0]
    lang_default.set(docx.oxml.shared.qn('w:val'), 'pt-BR')

    obj_charstyle = doc.styles.add_style('CommentsStyle', WD_STYLE_TYPE.PARAGRAPH)
    obj_font = obj_charstyle.font
    obj_font.size = Pt(13)
    obj_font.name = 'Calibri Light'

    doc.add_heading('Relatório Analítico', 0)
    doc.add_heading('Visão Clientes', 1)
    doc.paragraphs[1].runs[0].add_break(docx.enum.text.WD_BREAK.LINE)

    for k1 in relatorios.keys():
        for k2 in relatorios[k1].keys():
            doc.add_heading(k2, 2)
            if k2 != 'Campanha Table':
                doc.add_picture(
                    os.path.join('.', 'images', k2.lower().replace(' ', '_') + '.png'),
                    width=docx.shared.Cm(14.79), height=docx.shared.Cm(8)
                )
            else:
                rows = len(figures[k1][k2]['Columns'][0])
                cols = len(figures[k1][k2]['Columns'])
                table = doc.add_table(rows=rows, cols=cols)

                for i in range(rows):
                    row = table.rows[i].cells
                    for n, j in enumerate(figures[k1][k2]['Columns']):
                        row[n].text = j[i]

                for i in figures[k1][k2]['Values']:
                    row = table.add_row().cells
                    for n, j in enumerate(i):
                        row[n].text = str(j)

            if relatorios[k1][k2] is not None:
                doc.add_paragraph(relatorios[k1][k2], style='CommentsStyle')

    for arquivo in os.listdir(os.path.join('.', 'images')):
        if os.path.splitext(arquivo)[1].lower() in ['.jpg', '.jpeg', '.png']:
            os.remove(os.path.join('.', 'images', arquivo))

    return doc
